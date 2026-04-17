import csv
import re
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.db.models import (
    CaptureMediaORM,
    CaptureRecordORM,
    InspectionEvidenceORM,
    InspectionResultORM,
    InspectionTaskORM,
    ProjectORM,
    ProjectStructureInstanceORM,
)
from app.models.enums import StructureObjectType
from app.services.project_structures import get_object_meta, get_part_name, list_structure_part_templates
from app.services.template_store import get_enabled_inspection_items


def _ensure_export_dir(storage_root: Path, task_id: str) -> Path:
    export_dir = storage_root / "exports" / task_id
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def _now_suffix() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _item_meta_map(task: InspectionTaskORM) -> dict[str, dict[str, str]]:
    items = get_enabled_inspection_items(task.dam_type, task.enabled_chapters)
    return {
        item["item_code"]: {
            "chapter_code": item["chapter_code"],
            "item_name": item["item_name"],
        }
        for item in items
    }


def _sanitize_filename_part(value: str, max_length: int = 24) -> str:
    text = re.sub(r"[\\/:*?\"<>|\r\n]+", "_", value.strip())
    text = re.sub(r"\s+", "", text)
    if not text:
        return "未命名"
    return text[:max_length]


def _short_note(capture: CaptureRecordORM, max_length: int = 18) -> str:
    note = (capture.speech_text or capture.raw_note or "").strip()
    if not note:
        return "无描述"
    note = re.sub(r"\s+", "", note)
    return note[:max_length]


def _resolve_storage_file(storage_root: Path, server_url: str | None, local_path: str | None) -> Path | None:
    if local_path:
        path = Path(local_path)
        if path.exists():
            return path
    if server_url:
        cleaned = server_url.strip()
        if cleaned.startswith("/storage/"):
            path = storage_root / cleaned.removeprefix("/storage/")
            if path.exists():
                return path
    return None


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)


def _set_run_font(run, *, bold: bool = False, size_pt: int = 12) -> None:
    run.font.name = "Times New Roman"
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def _set_paragraph_style(paragraph, *, bold: bool = False, size_pt: int = 12, align=None) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, bold=bold, size_pt=size_pt)


def _write_cell_text(cell, text: str, *, bold: bool = False, size_pt: int = 12, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, bold=bold, size_pt=size_pt)
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_after = Pt(0)
    fmt.space_before = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _format_basic_value(value) -> str:
    if value is None:
        return ""
    return str(value)


def _collect_confirmed_captures(db: Session, task_id: str) -> list[CaptureRecordORM]:
    return db.scalars(
        select(CaptureRecordORM).where(
            CaptureRecordORM.task_id == task_id,
            CaptureRecordORM.deleted_at.is_(None),
            CaptureRecordORM.review_status == "confirmed",
        )
    ).all()


def export_issue_list_csv(db: Session, task: InspectionTaskORM, storage_root: Path) -> tuple[str, str]:
    rows = db.query(InspectionResultORM).filter(
        InspectionResultORM.task_id == task.task_id,
        InspectionResultORM.deleted_at.is_(None),
        InspectionResultORM.issue_flag.is_(True),
    ).all()

    meta_map = _item_meta_map(task)
    export_dir = _ensure_export_dir(storage_root, task.task_id)
    file_name = f"issue_list_{task.task_id}_{_now_suffix()}.csv"
    file_path = export_dir / file_name

    with file_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(
            [
                "chapter_code",
                "item_code",
                "item_name",
                "issue_type",
                "severity_level",
                "check_record",
                "suggestion",
            ]
        )
        for row in rows:
            meta = meta_map.get(row.item_code, {})
            writer.writerow(
                [
                    meta.get("chapter_code", ""),
                    row.item_code,
                    meta.get("item_name", row.item_code),
                    ",".join(row.issue_type or []),
                    row.severity_level or "",
                    row.check_record or "",
                    row.suggestion or "",
                ]
            )

    file_url = f"/storage/exports/{task.task_id}/{file_name}"
    return file_name, file_url


def export_photo_sheet_csv(db: Session, task: InspectionTaskORM, storage_root: Path) -> tuple[str, str]:
    results = db.query(InspectionResultORM).filter(
        InspectionResultORM.task_id == task.task_id,
        InspectionResultORM.deleted_at.is_(None),
    ).all()
    result_map = {r.result_id: r for r in results}
    meta_map = _item_meta_map(task)

    evidence_rows = db.query(InspectionEvidenceORM).filter(
        InspectionEvidenceORM.result_id.in_(list(result_map.keys()) if result_map else [""]),
        InspectionEvidenceORM.deleted_at.is_(None),
        InspectionEvidenceORM.evidence_type == "photo",
    ).all()

    export_dir = _ensure_export_dir(storage_root, task.task_id)
    file_name = f"photo_sheet_{task.task_id}_{_now_suffix()}.csv"
    file_path = export_dir / file_name

    with file_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(
            [
                "chapter_code",
                "item_code",
                "item_name",
                "evidence_id",
                "evidence_type",
                "caption",
                "shot_time",
                "file_url",
            ]
        )
        for evi in evidence_rows:
            result = result_map.get(evi.result_id)
            if result is None:
                continue
            meta = meta_map.get(result.item_code, {})
            writer.writerow(
                [
                    meta.get("chapter_code", ""),
                    result.item_code,
                    meta.get("item_name", result.item_code),
                    evi.evidence_id,
                    evi.evidence_type,
                    evi.caption or "",
                    evi.shot_time.isoformat() if evi.shot_time else "",
                    evi.file_url,
                ]
            )

    file_url = f"/storage/exports/{task.task_id}/{file_name}"
    return file_name, file_url


def export_photo_package_zip(
    db: Session,
    task: InspectionTaskORM,
    storage_root: Path,
) -> tuple[str, str]:
    structure_rows = db.scalars(
        select(ProjectStructureInstanceORM).where(
            ProjectStructureInstanceORM.project_id == task.project_id,
            ProjectStructureInstanceORM.deleted_at.is_(None),
            ProjectStructureInstanceORM.enabled_for_report.is_(True),
        )
    ).all()
    structure_map = {row.instance_id: row for row in structure_rows}
    captures = _collect_confirmed_captures(db, task.task_id)
    media_rows = db.scalars(
        select(CaptureMediaORM).where(
            CaptureMediaORM.capture_id.in_([row.capture_id for row in captures] or [""]),
            CaptureMediaORM.deleted_at.is_(None),
            CaptureMediaORM.media_type == "photo",
        )
    ).all()
    media_by_capture: dict[str, list[CaptureMediaORM]] = defaultdict(list)
    for media in media_rows:
        media_by_capture[media.capture_id].append(media)

    export_dir = _ensure_export_dir(storage_root, task.task_id)
    file_name = f"照片打包导出_{task.task_id}_{_now_suffix()}.zip"
    file_path = export_dir / file_name

    exported_count = 0
    with zipfile.ZipFile(file_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for capture in captures:
            structure = structure_map.get(capture.structure_instance_id)
            if structure is None:
                continue
            part_name = get_part_name(structure.template_source_type, capture.part_code) or capture.part_code
            folder_name = _sanitize_filename_part(structure.instance_name, max_length=40)
            for index, media in enumerate(media_by_capture.get(capture.capture_id, []), start=1):
                source = _resolve_storage_file(storage_root, media.server_url, media.local_path)
                if source is None:
                    continue
                ext = source.suffix or ".jpg"
                file_stub = "_".join(
                    [
                        _sanitize_filename_part(structure.instance_name, max_length=18),
                        _sanitize_filename_part(part_name, max_length=18),
                        _sanitize_filename_part(_short_note(capture), max_length=18),
                    ]
                )
                if len(media_by_capture.get(capture.capture_id, [])) > 1:
                    file_stub = f"{file_stub}_{index}"
                zf.write(source, arcname=f"{folder_name}/{file_stub}{ext}")
                exported_count += 1

        if exported_count == 0:
            zf.writestr(
                "导出说明.txt",
                "当前任务还没有已归档照片，或照片文件无法读取，因此没有可导出的图片。",
            )

    file_url = f"/storage/exports/{task.task_id}/{file_name}"
    return file_name, file_url


def export_inspection_docx(
    db: Session,
    task: InspectionTaskORM,
    storage_root: Path,
) -> tuple[str, str]:
    project = db.scalar(
        select(ProjectORM).where(
            ProjectORM.project_id == task.project_id,
            ProjectORM.deleted_at.is_(None),
        )
    )
    if project is None:
        raise ValueError("project not found")

    structure_rows = db.scalars(
        select(ProjectStructureInstanceORM).where(
            ProjectStructureInstanceORM.project_id == task.project_id,
            ProjectStructureInstanceORM.deleted_at.is_(None),
            ProjectStructureInstanceORM.enabled_for_report.is_(True),
        ).order_by(ProjectStructureInstanceORM.sort_order.asc())
    ).all()
    confirmed_captures = _collect_confirmed_captures(db, task.task_id)
    capture_notes: dict[tuple[str, str], list[str]] = defaultdict(list)
    abnormal_notes: list[str] = []
    for capture in confirmed_captures:
        note = (capture.speech_text or capture.raw_note or "").strip()
        if note:
            capture_notes[(capture.structure_instance_id, capture.part_code)].append(note)
            if capture.quick_status == "abnormal":
                abnormal_notes.append(note)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = document.add_paragraph()
    title_run = title.add_run("附录A 现场安全检查表")
    _set_run_font(title_run, bold=True, size_pt=16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("A.1 现场安全检查基本情况")
    _set_run_font(subtitle_run, bold=True, size_pt=14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    basic_table = document.add_table(rows=0, cols=2)
    basic_table.style = "Table Grid"
    basic_rows = [
        ("水库名称及基本情况描述", project.reservoir_name),
        ("枢纽工程主要建筑物", task.hub_main_structures or ""),
        ("水库防洪保护对象", task.flood_protect_obj or ""),
        ("检查时间", str(task.inspection_date)),
        ("天气", task.weather or ""),
        ("检查时库水位/m", _format_basic_value(task.water_level)),
        ("检查时库容/m³", _format_basic_value(task.storage)),
        ("检查人员", "、".join(task.inspectors or [])),
        ("现场检查发现的主要问题描述", "；".join(abnormal_notes[:8]) if abnormal_notes else "无"),
    ]
    for label, value in basic_rows:
        row = basic_table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.8)
        _write_cell_text(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_text(row.cells[1], value or "", align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_width(row.cells[0], 6.2)
        _set_cell_width(row.cells[1], 6.2)

    note_para = document.add_paragraph("注：可根据工程实际情况增减表中内容。")
    _set_paragraph_style(note_para, size_pt=12)

    for index, structure in enumerate(structure_rows, start=2):
        # The instance carries user-defined name, so the heading uses the actual configured object name.
        heading = document.add_paragraph()
        heading_run = heading.add_run(f"A.{index} {structure.instance_name}现场检查情况")
        _set_run_font(heading_run, bold=True, size_pt=14)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        header = table.rows[0]
        header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        header.height = Cm(0.8)
        _write_cell_text(header.cells[0], "对象实例", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_text(header.cells[1], "检查部位", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_text(header.cells[2], "检查情况记录", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_width(header.cells[0], 3.0)
        _set_cell_width(header.cells[1], 5.5)
        _set_cell_width(header.cells[2], 7.5)

        part_templates = list_structure_part_templates(StructureObjectType(structure.template_source_type))
        for part in part_templates:
            row = table.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Cm(0.8)
            part_code = str(part["part_code"])
            part_name = str(part["part_name"])
            notes = capture_notes.get((structure.instance_id, part_code), [])
            record_text = "；".join(notes) if notes else "未检查"
            _write_cell_text(row.cells[0], structure.instance_name, align=WD_ALIGN_PARAGRAPH.CENTER)
            _write_cell_text(row.cells[1], part_name, align=WD_ALIGN_PARAGRAPH.CENTER)
            _write_cell_text(row.cells[2], record_text, align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_width(row.cells[0], 3.0)
            _set_cell_width(row.cells[1], 5.5)
            _set_cell_width(row.cells[2], 7.5)

        if index != len(structure_rows) + 1:
            document.add_paragraph("")

    export_dir = _ensure_export_dir(storage_root, task.task_id)
    file_name = f"检查表格导出_{task.task_id}_{_now_suffix()}.docx"
    file_path = export_dir / file_name
    document.save(file_path)
    file_url = f"/storage/exports/{task.task_id}/{file_name}"
    return file_name, file_url
